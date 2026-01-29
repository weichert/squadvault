from __future__ import annotations

import sys
from pathlib import Path
from docx import Document

REPO_ROOT = Path(__file__).resolve().parent.parent
OPS_DIR = REPO_ROOT / "docs/30_contract_cards/ops"

SENTINEL = "[SV_CANONICAL_HEADER_V1]"

FILES = [
    ("Tone Engine", "CANONICAL", "v1.0", OPS_DIR / "TONE_ENGINE_Contract_Card_v1.0.docx"),
    ("Approval Authority", "CANONICAL", "v1.0", OPS_DIR / "APPROVAL_AUTHORITY_Contract_Card_v1.0.docx"),
    ("Version Presentation & Navigation", "CANONICAL", "v1.0", OPS_DIR / "VERSION_PRESENTATION_NAVIGATION_Contract_Card_v1.0.docx"),
    ("EAL Calibration", "CANONICAL", "v1.0", OPS_DIR / "EAL_CALIBRATION_Contract_Card_v1.0.docx"),
    ("Contract Validation Strategy", "CANONICAL", "v1.0", OPS_DIR / "CONTRACT_VALIDATION_STRATEGY_Contract_Card_v1.0.docx"),
]

def doc_has_sentinel(doc: Document) -> bool:
    s = SENTINEL.lower()
    return any((p.text or "").lower().find(s) >= 0 for p in doc.paragraphs[:30])

def prepend_header(doc: Document, name: str, status: str, version: str) -> None:
    # Insert at very top, in reverse order so final ordering is correct
    p0 = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph("")
    lines = [
        SENTINEL,
        f"Contract Name: {name}",
        f"Status: {status}",
        f"Version: {version}",
        "Authority: Tier 2 (Operational Contract Card)",
        "Canonical Location: docs/30_contract_cards/ops/",
        "",
    ]
    cur = p0
    # We want the block *above* existing content; insert in reverse after a dummy,
    # then swap by inserting before first paragraph isn't supported; instead:
    # create a new paragraph at start by inserting after, then move text upward by order.
    # python-docx doesn't support true insert-before; the stable pattern is:
    # set first paragraph text to first line, then insert subsequent lines after it.
    cur.text = lines[0]
    for line in lines[1:]:
        cur = cur.insert_paragraph_after(line)

def main() -> int:
    missing = [str(p) for _, _, _, p in FILES if not p.exists()]
    if missing:
        print("ERROR: missing expected contract card(s):", file=sys.stderr)
        for m in missing:
            print(f"  - {m}", file=sys.stderr)
        return 2

    failed = []
    for name, status, version, path in FILES:
        try:
            doc = Document(str(path))
        except Exception as e:
            failed.append((str(path), f"open failed: {e}"))
            continue

        if doc_has_sentinel(doc):
            print(f"SKIP: sentinel present: {path}")
            continue

        prepend_header(doc, name=name, status=status, version=version)
        try:
            doc.save(str(path))
            print(f"OK: prepended canonical header: {path}")
        except Exception as e:
            failed.append((str(path), f"save failed: {e}"))

    if failed:
        print("\nFAIL: one or more contract cards could not be patched (fail-closed).", file=sys.stderr)
        for p, why in failed:
            print(f"  - {p}: {why}", file=sys.stderr)
        return 2

    return 0

if __name__ == "__main__":
    raise SystemExit(main())
